#!/usr/bin/env python3
"""
投委会备忘录 .docx 生成器
用法: python generate-memo-docx.py --data memo-data.json --output 备忘录.docx

依赖: pip install python-docx
"""

import json
import sys
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print("请先安装: pip install python-docx")
    sys.exit(1)

# 配色（深蓝+金——专业投资风格）
C = {
    "navy": RGBColor(0x0D, 0x2B, 0x4E),
    "gold": RGBColor(0xC8, 0x96, 0x3E),
    "white": RGBColor(0xFF, 0xFF, 0xFF),
    "stripe": RGBColor(0xF2, 0xF5, 0xF9),
    "red": RGBColor(0xC0, 0x39, 0x2B),
    "green": RGBColor(0x27, 0xAE, 0x60),
    "gray": RGBColor(0x99, 0x99, 0x99),
}


def shade(cell, hex_color):
    """设置单元格背景色"""
    cell._tc.get_or_add_tcPr().append(
        parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>'))


def styled_table(doc, headers, rows, col_widths=None):
    """深蓝表头 + 斑马纹表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = C["white"]
        run.font.name = "Calibri"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        shade(cell, "0D2B4E")

    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val) if val is not None else "")
            run.font.size = Pt(9.5)
            run.font.name = "Calibri"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            if r % 2 == 0:
                shade(cell, "F2F5F9")

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph("")
    return table


def build_memo(data: dict, output_path: str):
    doc = Document()

    # 页面设置 A4
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)

    # 默认字体
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.35

    # ═══ 封面 ═══
    for _ in range(6):
        doc.add_paragraph("")

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("投资备忘录")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = C["navy"]
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    doc.add_paragraph("")
    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = line.add_run("━" * 30)
    r.font.color.rgb = C["gold"]
    r.font.size = Pt(10)
    doc.add_paragraph("")

    for label, val in [
        ("项目名称", data.get("company", "【项目名称】")),
        ("日期", data.get("date", datetime.now().strftime("%Y年%m月%d日"))),
        ("分析师", data.get("analyst", "【分析师】")),
        ("保密级别", "绝密——仅限投委会内部"),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"{label}：{val}")
        r.font.size = Pt(12)
        r.font.color.rgb = C["navy"]
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.add_page_break()

    # ═══ 章节辅助函数 ═══
    def h1(text):
        h = doc.add_heading(text, level=1)
        for r in h.runs:
            r.font.color.rgb = C["navy"]
            r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    # ═══ 1. 执行摘要 ═══
    h1("1. 执行摘要")
    doc.add_paragraph(data.get("executive_summary", "【在此填写——≤200字】"))
    verdict = data.get("verdict", "待定")
    p = doc.add_paragraph()
    r = p.add_run(f"建议：{verdict}")
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = C["green"] if "投" in verdict else C["red"]

    # ═══ 2. 公司概览 ═══
    h1("2. 公司概览")
    styled_table(doc, ["维度", "详情"],
        [[k, data.get(k, "")] for k in [
            "公司全称", "成立时间", "总部", "赛道", "阶段",
            "本轮融资金额", "估值", "核心产品", "员工数"
        ]], col_widths=[4, 12])

    # ═══ 3. 投资论点 ═══
    h1("3. 投资论点")
    doc.add_heading("3.1 为什么投？", level=2)
    for i, t in enumerate(data.get("thesis", []), 1):
        doc.add_paragraph(f"{i}. {t}", style='List Number')

    doc.add_heading("3.2 关键假设与可证伪条件", level=2)
    assumptions = data.get("key_assumptions", [])
    if assumptions:
        styled_table(doc, ["假设", "如果错了会怎样", "验证方法"],
            [[a.get("assumption", ""), a.get("if_wrong", ""), a.get("verify", "")]
             for a in assumptions], col_widths=[5, 5.5, 5.5])

    # ═══ 4. 风险 ═══
    h1("4. 风险与缓解")
    risks = data.get("risks", [])
    if risks:
        styled_table(doc, ["风险", "严重度", "概率", "缓解措施"],
            [[r.get("risk", ""), r.get("severity", ""), r.get("probability", ""),
              r.get("mitigation", "")] for r in risks], col_widths=[6, 2, 2, 6])

    doc.add_heading("4.1 致命风险（一旦发生归零）", level=2)
    for r in data.get("fatal_risks", []):
        doc.add_paragraph(f"▪ {r}", style='List Bullet')

    # ═══ 5. 市场与竞争 ═══
    h1("5. 市场与竞争")
    doc.add_heading("5.1 市场规模", level=2)
    styled_table(doc, ["指标", "数值", "来源"],
        [[k, data.get(k, ""), data.get(f"{k}_source", "")] for k in
         ["TAM", "SAM", "SOM", "CAGR"]], col_widths=[5, 5, 6])

    doc.add_heading("5.2 竞争格局", level=2)
    comps = data.get("competitors", [])
    comp_rows = [[data.get("company", "标的"), data.get("stage", ""),
        data.get("valuation", ""), data.get("differentiation", ""),
        data.get("total_funding", "")]]
    for c in comps:
        comp_rows.append([c.get("name", ""), c.get("stage", ""),
            c.get("valuation", ""), c.get("differentiation", ""),
            c.get("total_funding", "")])
    styled_table(doc, ["公司", "阶段", "估值", "差异化", "融资总额"],
        comp_rows, col_widths=[3, 2, 2.5, 5, 3.5])

    # ═══ 6. 团队 ═══
    h1("6. 团队评估")
    doc.add_heading("6.1 创始人", level=2)
    founder = data.get("founder", {})
    styled_table(doc, ["维度", "详情"],
        [[k, founder.get(k, "")] for k in
         ["姓名", "教育", "此前经历", "行业经验", "决策模式"]],
        col_widths=[4, 12])

    doc.add_heading("6.2 核心团队", level=2)
    team = data.get("team", [])
    if team:
        styled_table(doc, ["角色", "姓名", "背景", "评价"],
            [[t.get("role", ""), t.get("name", ""), t.get("background", ""),
              t.get("assessment", "")] for t in team],
            col_widths=[3, 2.5, 6, 4.5])

    # ═══ 7. 财务 ═══
    h1("7. 财务与回报")
    fin = data.get("financials", {})
    styled_table(doc, ["指标", "去年", "今年(预计)", "明年(预计)"],
        [["营收", fin.get("revenue_ly", ""), fin.get("revenue_ey", ""), fin.get("revenue_ny", "")],
         ["增速", fin.get("growth_ly", ""), fin.get("growth_ey", ""), fin.get("growth_ny", "")],
         ["毛利率", fin.get("gm_ly", ""), fin.get("gm_ey", ""), fin.get("gm_ny", "")],
         ["月烧钱", fin.get("burn_ly", ""), fin.get("burn_ey", ""), fin.get("burn_ny", "")]],
        col_widths=[5, 3.7, 3.7, 3.7])

    doc.add_heading("7.1 回报情景", level=2)
    scenarios = data.get("scenarios", [
        {"name": "乐观", "exit_val": "", "moic": "", "irr": "", "probability": ""},
        {"name": "基准", "exit_val": "", "moic": "", "irr": "", "probability": ""},
        {"name": "悲观", "exit_val": "", "moic": "", "irr": "", "probability": ""},
    ])
    styled_table(doc, ["情景", "退出估值", "MOIC", "IRR", "概率"],
        [[s.get("name", ""), s.get("exit_val", ""), s.get("moic", ""),
          s.get("irr", ""), s.get("probability", "")] for s in scenarios],
        col_widths=[3, 3.5, 3, 3, 3.5])

    # ═══ 8. 尽调发现 ═══
    h1("8. 尽调关键发现")
    dd = data.get("due_diligence", {})
    for label, key in [("洞察视角", "insight"), ("审计视角", "audit"), ("工程视角", "engineering")]:
        p = doc.add_paragraph()
        r = p.add_run(f"【{label}】")
        r.bold = True
        p.add_run(f" {dd.get(key, '【待填写】')}")

    disagreements = data.get("disagreements", [])
    if disagreements:
        doc.add_heading("8.1 视角分歧", level=2)
        styled_table(doc, ["议题", "分歧", "谁证据更强"],
            [[d.get("topic", ""), d.get("disagreement", ""), d.get("stronger", "")]
             for d in disagreements], col_widths=[4, 7, 5])

    # ═══ 9. 条款 ═══
    h1("9. 条款摘要")
    terms = data.get("terms", {})
    styled_table(doc, ["条款", "详情"],
        [[k, terms.get(k, "")] for k in [
            "投资金额", "估值", "股权比例", "领投/跟投", "董事会席位",
            "一票否决权", "反稀释", "清算优先权", "其他特殊条款"
        ]], col_widths=[4, 12])

    # ═══ 10. 建议 ═══
    h1("10. 建议与下一步")
    doc.add_paragraph(data.get("recommendation", "【待填写】"))
    next_steps = data.get("next_steps", [])
    if next_steps:
        styled_table(doc, ["动作", "负责人", "截止日期"],
            [[s.get("action", ""), s.get("owner", ""), s.get("deadline", "")]
             for s in next_steps], col_widths=[9, 3.5, 3.5])

    # ═══ 免责声明 ═══
    doc.add_paragraph("")
    p = doc.add_paragraph()
    r = p.add_run("免责声明：本备忘录为内部分析文件，不构成投资建议。")
    r.font.size = Pt(8)
    r.font.color.rgb = C["gray"]
    r.italic = True

    # ═══ 页眉页脚 ═══
    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = hp.add_run("绝密——仅限投委会内部")
        r.font.size = Pt(8)
        r.font.color.rgb = C["red"]

        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = fp.add_run("— 本地AI自动生成 —")
        r.font.size = Pt(8)
        r.font.color.rgb = C["gray"]

    doc.save(output_path)
    print(f"备忘录已生成: {output_path}")


if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser(description="生成投委会备忘录 .docx")
    parser.add_argument("--data", required=True, help="JSON 数据文件")
    parser.add_argument("--output", default="投资备忘录.docx", help="输出路径")
    args = parser.parse_args()
    with open(args.data, "r", encoding="utf-8") as f:
        build_memo(json.load(f), args.output)
